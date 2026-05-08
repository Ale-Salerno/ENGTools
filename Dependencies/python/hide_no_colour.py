import os
import webbrowser
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Manual mapping of highlight color indices to RGB values and names
HIGHLIGHT_COLOR_MAP = {
    1: ('#000000', 'Black'),
    2: ('#00008B', 'Dark Blue'),
    3: ('#00FFFF', 'Cyan'),
    4: ('#00FF00', 'Bright Green'),
    5: ('#40E0D0', 'Turquoise'),
    6: ('#FF0000', 'Red'),
    7: ('#FFFF00', 'Yellow'),
    8: ('#808080', 'Gray 50%'),
    9: ('#008000', 'Green'),
    10: ('#FFC0CB', 'Pink'),
    11: ('#808000', 'Dark Yellow'),
    12: ('#008080', 'Teal'),
    13: ('#8B0000', 'Dark Red'),
    14: ('#EE82EE', 'Violet'),
    15: ('#FFFFFF', 'White'),
    16: ('#D3D3D3', 'Light Gray'),
}

def gather_font_colors(doc):
    """Collect unique font colors from the document."""
    colors = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip() and run.font.color and run.font.color.rgb:
                colors.add(str(run.font.color.rgb).upper())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip() and run.font.color and run.font.color.rgb:
                            colors.add(str(run.font.color.rgb).upper())
    return colors

def gather_highlight_colors(doc):
    """Collect unique highlight colors from the document."""
    highlights = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text.strip() and run.font.highlight_color not in [None, 0]:
                highlights.add(run.font.highlight_color)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text.strip() and run.font.highlight_color not in [None, 0]:
                            highlights.add(run.font.highlight_color)
    return highlights

def gather_cell_shading_colors(doc):
    """Collect unique cell shading colors from the document's tables."""
    shading_colors = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._element.tcPr
                shd = tcPr.find(qn("w:shd"))
                if shd is not None:
                    fill = shd.get(qn("w:fill"))
                    if fill:
                        hex_color = "#" + fill.upper()
                        shading_colors.add(hex_color)
    return shading_colors

def generate_html_from_colors(font_colors, highlight_colors, shading_colors, html_path):
    """Generate HTML with color tables."""
    sorted_font = sorted(font_colors)
    sorted_high = sorted(highlight_colors)
    sorted_shading = sorted(shading_colors)
    
    html = [
        "<html><head><meta charset='UTF-8'><title>Color Samples</title>",
        "<style>table, th, td { border: 1px solid black; border-collapse: collapse; padding: 5px; }",
        ".swatch { width: 50px; height: 20px; display: inline-block; border: 1px solid #000; }",
        "</style></head><body>",
        "<h2>Font Colors</h2><table><tr><th>Index</th><th>Hex (RGB)</th><th>Swatch</th></tr>"
    ]
    
    # Font colors table
    for idx, color in enumerate(sorted_font, 1):
        html.append(
            f"<tr><td>{idx}</td><td>#{color}</td>"
            f"<td><div class='swatch' style='background-color:#{color};'></div></td></tr>"
        )
    html.append("</table>")
    
    # Highlight colors table
    html.append("<h2>Highlight Colors</h2><table><tr><th>Index</th><th>Color</th><th>Swatch</th></tr>")
    start_high_idx = len(sorted_font) + 1
    for i, hl in enumerate(sorted_high, start_high_idx):
        rgb, name = HIGHLIGHT_COLOR_MAP.get(hl, ('#FFFFFF', 'Unknown'))
        html.append(
            f"<tr><td>{i}</td><td>{name}</td>"
            f"<td><div class='swatch' style='background-color:{rgb};'></div></td></tr>"
        )
    html.append("</table>")
    
    # Cell shading colors table
    html.append("<h2>Cell Shading Colors</h2><table><tr><th>Index</th><th>Hex (RGB)</th><th>Swatch</th></tr>")
    start_shading_idx = len(sorted_font) + len(sorted_high) + 1
    for idx, color in enumerate(sorted_shading, start_shading_idx):
        html.append(
            f"<tr><td>{idx}</td><td>{color}</td>"
            f"<td><div class='swatch' style='background-color:{color};'></div></td></tr>"
        )
    html.append("</table></body></html>")
    
    with open(html_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(html))
    return sorted_font, sorted_high, sorted_shading

def prompt_for_indices(prompt_text, max_index):
    """Prompt user for indices and return as a set."""
    chosen = set()
    inp = input(prompt_text).strip()
    if not inp:
        return chosen
    for part in inp.split(','):
        part = part.strip()
        if part.isdigit():
            num = int(part)
            if 1 <= num <= max_index:
                chosen.add(num)
    return chosen

def should_hide_run(run, font_colors, highlights, mode):
    """Determine if a run should be hidden based on font and highlight colors."""
    font_color = str(run.font.color.rgb).upper() if run.font.color and run.font.color.rgb else None
    highlight_color = run.font.highlight_color
    
    font_selected = bool(font_colors)
    highlight_selected = bool(highlights)
    
    if mode == "hide_not_chosen":
        font_cond = font_selected and (font_color not in font_colors)
        highlight_cond = highlight_selected and (highlight_color not in highlights)
        return font_cond or highlight_cond
    elif mode == "hide_chosen":
        font_cond = font_selected and (font_color in font_colors)
        highlight_cond = highlight_selected and (highlight_color in highlights)
        return font_cond or highlight_cond
    return False

def set_run_hidden(run, hidden=True):
    """Toggle the hidden state of a run."""
    rPr = run._element.find(qn("w:rPr"))
    if not rPr:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)
    vanish = rPr.find(qn("w:vanish"))
    if hidden:
        if not vanish:
            vanish = OxmlElement("w:vanish")
            vanish.set(qn("w:val"), "true")
            rPr.append(vanish)
    elif vanish is not None:
        rPr.remove(vanish)

def get_cell_shading_color(cell):
    """Get the shading color of a cell as a hex string, or None if not set."""
    tcPr = cell._element.tcPr
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        fill = shd.get(qn("w:fill"))
        if fill:
            return "#" + fill.upper()
    return None

def process_paragraphs(paragraphs, font_colors, highlights, mode):
    """Process runs in paragraphs."""
    for para in paragraphs:
        for run in para.runs:
            hide = should_hide_run(run, font_colors, highlights, mode)
            set_run_hidden(run, hide)

def hide_text_in_doc(doc, font_colors, highlights, selected_shading, mode):
    """Apply hiding to document elements."""
    process_paragraphs(doc.paragraphs, font_colors, highlights, mode)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs, font_colors, highlights, mode)
                # Check cell shading color
                cell_shading_color = get_cell_shading_color(cell)
                if cell_shading_color is not None:
                    shading_selected = bool(selected_shading)
                    if mode == "hide_not_chosen":
                        if shading_selected and (cell_shading_color not in selected_shading):
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    set_run_hidden(run, True)
                    elif mode == "hide_chosen":
                        if shading_selected and (cell_shading_color in selected_shading):
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    set_run_hidden(run, True)

def process_docx_file(file_path, font_colors, highlights, selected_shading, mode):
    """Process a DOCX file."""
    try:
        doc = Document(file_path)
        hide_text_in_doc(doc, font_colors, highlights, selected_shading, mode)
        doc.save(file_path)
        print(f"[OK] Processed: {file_path}")
    except Exception as e:
        print(f"[ERROR] Processing {file_path}: {e}")

def main():
    cwd = os.getcwd()
    docx_files = [f for f in os.listdir(cwd) if f.lower().endswith(".docx")]
    if not docx_files:
        print("No DOCX files found.")
        return

    print("Gathering colors from DOCX files...")
    all_font_colors = set()
    all_highlights = set()
    all_shading = set()
    for filename in docx_files:
        path = os.path.join(cwd, filename)
        try:
            doc = Document(path)
            all_font_colors.update(gather_font_colors(doc))
            all_highlights.update(gather_highlight_colors(doc))
            all_shading.update(gather_cell_shading_colors(doc))
            print(f" - {filename}: Font={len(all_font_colors)}, Highlights={len(all_highlights)}, Shading={len(all_shading)}")
        except Exception as e:
            print(f"[ERROR] {filename}: {e}")

    html_path = os.path.join(cwd, "color_samples.html")
    sorted_font, sorted_high, sorted_shading = generate_html_from_colors(all_font_colors, all_highlights, all_shading, html_path)
    print(f"\nHTML generated: {html_path}")
    webbrowser.open("file://" + html_path)

    max_index = len(sorted_font) + len(sorted_high) + len(sorted_shading)
    chosen_indices = prompt_for_indices(
        f"Enter indices (1-{max_index}, comma-separated) or Enter to skip: ",
        max_index
    )
    
    # Determine selected colors
    selected_font = set()
    selected_high = set()
    selected_shading = set()
    font_count = len(sorted_font)
    high_count = len(sorted_high)
    shading_count = len(sorted_shading)
    
    for idx in chosen_indices:
        if 1 <= idx <= font_count:
            selected_font.add(sorted_font[idx-1])
        elif font_count < idx <= font_count + high_count:
            hl_idx = idx - font_count - 1
            if 0 <= hl_idx < high_count:
                selected_high.add(sorted_high[hl_idx])
        else:
            shading_idx = idx - (font_count + high_count) - 1
            if 0 <= shading_idx < shading_count:
                selected_shading.add(sorted_shading[shading_idx])

    print("\nChoose mode:")
    print("1. Hide text NOT matching selected colors (default)")
    print("2. Hide text matching selected colors")
    mode = "hide_not_chosen" if input("Enter 1 or 2: ").strip() != '2' else "hide_chosen"

    print("\nProcessing files...")
    for filename in docx_files:
        path = os.path.join(cwd, filename)
        process_docx_file(path, selected_font, selected_high, selected_shading, mode)
    print("\nDone.")

if __name__ == "__main__":
    main()